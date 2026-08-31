from __future__ import annotations

import io
from dataclasses import dataclass, field
from pathlib import Path

import pandas as pd
from docx import Document as DocxDocument
from PIL import Image, ImageOps

MAX_TEXT_CHARS = 120_000
MAX_TABLE_ROWS = 4_000
MAX_IMAGE_EDGE = 2048


@dataclass
class ExtractedContent:
    text: str
    tables: list[pd.DataFrame] = field(default_factory=list)
    images: list[bytes] = field(default_factory=list)
    notes: list[str] = field(default_factory=list)


def _clip_text(value: str) -> str:
    if len(value) <= MAX_TEXT_CHARS:
        return value
    return value[:MAX_TEXT_CHARS] + "\n[truncated]"


def _sanitize_cell(value: object) -> object:
    if isinstance(value, str) and value[:1] in {"=", "+", "-", "@", "\t", "\r"}:
        return "'" + value
    return value


def _read_csv(path: Path) -> ExtractedContent:
    frame = pd.read_csv(path, nrows=MAX_TABLE_ROWS, dtype=str, keep_default_na=False)
    frame = frame.map(_sanitize_cell)
    return ExtractedContent(text=_clip_text(frame.to_csv(index=False)), tables=[frame])


def _read_xlsx(path: Path) -> ExtractedContent:
    sheets = pd.read_excel(path, sheet_name=None, nrows=MAX_TABLE_ROWS, dtype=str, engine="openpyxl")
    tables: list[pd.DataFrame] = []
    chunks: list[str] = []
    for name, frame in list(sheets.items())[:8]:
        frame = frame.fillna("").map(_sanitize_cell)
        tables.append(frame)
        chunks.append(f"# Sheet: {name}\n{frame.to_csv(index=False)}")
    return ExtractedContent(text=_clip_text("\n\n".join(chunks)), tables=tables)


def _read_docx(path: Path) -> ExtractedContent:
    document = DocxDocument(str(path))
    paragraphs = [p.text.strip() for p in document.paragraphs if p.text.strip()]
    tables: list[pd.DataFrame] = []
    for table in document.tables[:20]:
        rows = [[_sanitize_cell(cell.text.strip()) for cell in row.cells] for row in table.rows[:MAX_TABLE_ROWS]]
        if rows:
            header = [str(item) for item in rows[0]]
            body = rows[1:] if len(rows) > 1 else []
            tables.append(pd.DataFrame(body, columns=header))
    text = "\n".join(paragraphs)
    if tables:
        text += "\n\n" + "\n\n".join(table.to_csv(index=False) for table in tables)
    return ExtractedContent(text=_clip_text(text), tables=tables)


def _read_txt(path: Path) -> ExtractedContent:
    raw = path.read_bytes()[: MAX_TEXT_CHARS * 2]
    try:
        text = raw.decode("utf-8")
    except UnicodeDecodeError:
        text = raw.decode("latin-1")
    return ExtractedContent(text=_clip_text(text))


def _read_pdf(path: Path) -> ExtractedContent:
    try:
        import fitz
    except ImportError:
        return ExtractedContent(
            text="",
            notes=["PDF text extraction is unavailable in this environment."],
        )

    doc = fitz.open(path)
    try:
        if doc.page_count > 40:
            notes = [f"Only the first 40 of {doc.page_count} pages were analyzed."]
            page_count = 40
        else:
            notes = []
            page_count = doc.page_count

        texts: list[str] = []
        images: list[bytes] = []
        for index in range(page_count):
            page = doc.load_page(index)
            texts.append(page.get_text("text") or "")

        extracted = _clip_text("\n\n".join(texts)).strip()
        if len(extracted) < 200:
            for index in range(min(page_count, 5)):
                page = doc.load_page(index)
                pix = page.get_pixmap(matrix=fitz.Matrix(1.5, 1.5), alpha=False)
                images.append(pix.tobytes("png"))
            notes.append("Scanned PDF: page images were sent for visual review.")
        return ExtractedContent(text=extracted, images=images, notes=notes)
    finally:
        doc.close()


def _read_image(path: Path) -> ExtractedContent:
    with Image.open(path) as image:
        image = ImageOps.exif_transpose(image)
        if image.mode not in {"RGB", "L"}:
            image = image.convert("RGB")
        image.thumbnail((MAX_IMAGE_EDGE, MAX_IMAGE_EDGE))
        buffer = io.BytesIO()
        image.save(buffer, format="PNG", optimize=True)
        return ExtractedContent(text="", images=[buffer.getvalue()], notes=["Image document."])


def extract_document(path: Path, detected_type: str) -> ExtractedContent:
    readers = {
        "csv": _read_csv,
        "xlsx": _read_xlsx,
        "docx": _read_docx,
        "txt": _read_txt,
        "pdf": _read_pdf,
        "png": _read_image,
        "jpg": _read_image,
        "webp": _read_image,
        "tiff": _read_image,
    }
    reader = readers.get(detected_type)
    if reader is None:
        return ExtractedContent(text="", notes=["Unsupported type after save."])
    return reader(path)
